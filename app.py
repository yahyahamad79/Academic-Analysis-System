import streamlit as st
import pandas as pd
import numpy as np
import pingouin as pg
from docx import Document
from io import BytesIO
import plotly.express as px

# إعدادات الواجهة الأكاديمية
st.set_page_config(page_title="منصة الخبير الأكاديمي", layout="wide")

def create_word_report(ar_text, en_text):
    """توليد ملف Word احترافي باللغتين"""
    doc = Document()
    doc.add_heading('تقرير التحليل الإحصائي الأكاديمي (APA 7)', 0)
    
    doc.add_heading('أولاً: النتائج باللغة العربية', level=1)
    p_ar = doc.add_paragraph(ar_text)
    doc.add_page_break()
    
    doc.add_heading('Second: Academic Results (English)', level=1)
    doc.add_paragraph(en_text)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def main():
    st.title("🏛️ منصة المشرف الأكاديمي للتحليل الإحصائي المتقدم")
    st.write("---")

    file = st.file_uploader("ارفع ملف البيانات (Excel)", type=['xlsx'])
    
    if file:
        df = pd.read_excel(file)
        st.success("تم تحميل البيانات بنجاح.")
        
        tab1, tab2, tab3 = st.tabs(["📊 التحليل الشامل", "📝 تحليل الاستبيانات", "⚙️ فحص الفروض"])
        
        with tab1:
            if st.button("🚀 تشغيل التحليل الشامل وتوليد التقرير"):
                # 1. فحص التوزيع الطبيعي
                st.subheader("🔍 فحص التوزيع الطبيعي (Shapiro-Wilk)")
                norm = pg.normality(df.select_dtypes(include=[np.number]))
                st.table(norm)
                
                # 2. الإحصاء الوصفي
                st.subheader("📈 الإحصاء الوصفي")
                st.table(df.describe().T)
                
                # 3. مثال لتحليل فرضية (T-Test) مع حجم الأثر
                num_cols = df.select_dtypes(include=[np.number]).columns
                if len(num_cols) > 0:
                    res = pg.ttest(df[num_cols[0]], df[num_cols[0]].mean())
                    p_val = res['p-val'].values[0]
                    eff_size = res['cohen-d'].values[0]
                    
                    # صياغة التعقيب
                    ar_rep = f"أظهرت النتائج أن قيمة الدلالة بلغت ({p_val:.3f})، مما يشير إلى {'وجود' if p_val < 0.05 else 'عدم وجود'} أثر ذو دلالة إحصائية. كما بلغت قيمة حجم الأثر ({eff_size:.2f})."
                    en_rep = f"The analysis revealed a p-value of ({p_val:.3f}), indicating a {'statistically significant' if p_val < 0.05 else 'non-significant'} result. The effect size (Cohen's d) was ({eff_size:.2f})."
                    
                    st.subheader("📝 التقرير الأكاديمي المولد")
                    st.info(ar_rep)
                    st.info(en_rep)
                    
                    # تصدير الملف
                    word_data = create_word_report(ar_rep, en_rep)
                    st.download_button("📥 تحميل التقرير كاملاً (Word)", word_data, "Academic_Report.docx")

        with tab2:
            st.subheader("📊 تحليل محاور الاستبيان (النسب المئوية)")
            cols = st.multiselect("اختر فقرات المحور:", df.columns)
            if cols and st.button("تحليل التكرارات والتعقيب"):
                for c in cols:
                    st.write(f"**الفقرة: {c}**")
                    counts = df[c].value_counts(normalize=True)*100
                    st.bar_chart(counts)
                    st.write(f"المتوسط الحسابي: {df[c].mean():.2f}")

if __name__ == '__main__':
    main()
    